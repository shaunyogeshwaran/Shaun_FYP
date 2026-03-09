"""
Merge thesis_new_chapters.docx into IPD DOCUMENT.docx.

1. Replaces the old abstract (paragraphs 20-21) with the updated abstract
2. Inserts Chapters 6-8 between Chapter 05 and References
3. Saves as thesis_merged.docx (does not overwrite originals)

Run: python merge_documents.py
"""

import copy
from docx import Document
from docx.oxml.ns import qn


def find_paragraph_index(doc, match_fn):
    """Find the index of a paragraph matching a condition."""
    for i, p in enumerate(doc.paragraphs):
        if match_fn(p):
            return i
    return None


def get_body(doc):
    """Get the document body XML element."""
    return doc.element.body


def paragraph_element_at(doc, para_idx):
    """Get the underlying XML element for a paragraph by index."""
    return doc.paragraphs[para_idx]._element


def main():
    # Load documents
    print("Loading documents...")
    ipd = Document("IPD DOCUMENT.docx")
    new = Document("thesis_new_chapters.docx")

    ipd_body = get_body(ipd)

    # =========================================================================
    # Step 1: Replace the old abstract (paragraphs 20-21)
    # =========================================================================
    print("Replacing abstract...")

    # Find old abstract body paragraphs (between ABSTRACT heading and Subject Descriptors)
    abstract_heading_idx = find_paragraph_index(
        ipd, lambda p: p.text.strip() == "ABSTRACT" and p.style.name.startswith("Heading"))

    # Paragraphs 20 and 21 are the two abstract body paragraphs
    # Remove them
    old_abstract_p1 = paragraph_element_at(ipd, abstract_heading_idx + 1)
    old_abstract_p2 = paragraph_element_at(ipd, abstract_heading_idx + 2)

    # Find the abstract content in the new document (skip title page, get body paragraphs)
    new_abstract_paras = []
    found_abstract = False
    for p in new.paragraphs:
        if p.text.strip() == "Abstract" and p.style.name.startswith("Heading"):
            found_abstract = True
            continue
        if found_abstract:
            # Stop at the next chapter heading or page break
            if p.text.startswith("Chapter 6"):
                break
            if p.text.strip():  # Only non-empty paragraphs
                new_abstract_paras.append(p)

    print(f"  Found {len(new_abstract_paras)} new abstract paragraphs")

    # Insert new abstract paragraphs before the first old one, then remove old ones
    insert_before = old_abstract_p1
    for p in new_abstract_paras:
        new_elem = copy.deepcopy(p._element)
        insert_before.addprevious(new_elem)

    # Remove old abstract paragraphs
    ipd_body.remove(old_abstract_p1)
    ipd_body.remove(old_abstract_p2)

    # =========================================================================
    # Step 2: Insert Chapters 6-8 before References
    # =========================================================================
    print("Inserting Chapters 6-8...")

    # Re-find References heading (indices shifted after abstract edit)
    refs_idx = find_paragraph_index(
        ipd, lambda p: p.text.strip() == "References" and p.style.name.startswith("Heading"))
    refs_element = paragraph_element_at(ipd, refs_idx)

    # Also remove the empty Heading 2 paragraphs between Ch5 content and References
    # (they're just blank spacers)
    empty_to_remove = []
    for i in range(refs_idx - 1, 0, -1):
        p = ipd.paragraphs[i]
        if p.text.strip() == "" and p.style.name.startswith("Heading"):
            empty_to_remove.append(p._element)
        else:
            break

    for elem in empty_to_remove:
        ipd_body.remove(elem)

    # Re-find References after removing empties
    refs_idx = find_paragraph_index(
        ipd, lambda p: p.text.strip() == "References" and p.style.name.startswith("Heading"))
    refs_element = paragraph_element_at(ipd, refs_idx)

    # Collect all elements from new document for Chapters 6-8
    # (everything after the Abstract section)
    new_body = get_body(new)
    chapter_elements = []
    copying = False

    for elem in new_body:
        # Start copying from the page break before Chapter 6
        if not copying:
            # Look for Chapter 6 heading or page breaks leading to it
            if elem.tag == qn("w:p"):
                # Check if this paragraph contains "Chapter 6" text
                text = elem.text or ""
                # Also check runs
                for r in elem.findall(qn("w:r")):
                    t = r.find(qn("w:t"))
                    if t is not None and t.text:
                        text += t.text
                if "Chapter 6" in text:
                    copying = True
                    # Also include the preceding page break if it exists
            # Check for sectPr (page breaks can be in paragraph properties)
            if not copying:
                for pPr in elem.findall(qn("w:pPr")):
                    for secPr in pPr.findall(qn("w:sectPr")):
                        pass  # handled below

        if copying:
            # Skip the final sectPr (section properties at doc end)
            if elem.tag == qn("w:sectPr"):
                continue
            chapter_elements.append(elem)

    # Also need to grab page break before Chapter 6
    # Find it in the new doc elements
    all_new_elems = list(new_body)
    for i, elem in enumerate(all_new_elems):
        if elem.tag == qn("w:p"):
            text = ""
            for r in elem.findall(qn("w:r")):
                t = r.find(qn("w:t"))
                if t is not None and t.text:
                    text += t.text
            if "Chapter 6" in text:
                # Check if previous element is a page-break paragraph
                if i > 0:
                    prev = all_new_elems[i - 1]
                    if prev.tag == qn("w:p") and prev not in chapter_elements:
                        # Check for page break
                        has_break = False
                        for r in prev.findall(qn("w:r")):
                            for br in r.findall(qn("w:br")):
                                if br.get(qn("w:type")) == "page":
                                    has_break = True
                        if has_break:
                            chapter_elements.insert(0, prev)
                break

    print(f"  Collected {len(chapter_elements)} elements for Chapters 6-8")

    # Also copy tables from new document - they're separate from paragraphs in the body
    # Actually, tables are inline in the body element, so they should be in chapter_elements

    # Count tables in chapter_elements
    table_count = sum(1 for e in chapter_elements if e.tag == qn("w:tbl"))
    print(f"  Including {table_count} tables")

    # Insert all chapter elements before References
    for elem in chapter_elements:
        new_elem = copy.deepcopy(elem)
        refs_element.addprevious(new_elem)

    # Add a page break before References
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    page_break_xml = (
        f'<w:p {nsdecls("w")}>'
        f'<w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )
    refs_element.addprevious(parse_xml(page_break_xml))

    # =========================================================================
    # Step 3: Save merged document
    # =========================================================================
    output_path = "thesis_merged.docx"
    ipd.save(output_path)
    print(f"\nSaved merged document to {output_path}")


if __name__ == "__main__":
    main()
