#!/usr/bin/env python3
"""Format thesis_updated.docx to comply with IIT/Westminster FYP template.

Changes:
  1. Add Declaration and Acknowledgements pages (front matter)
  2. Restructure Abstract (3 paragraphs + Subject Descriptors + Keywords)
  3. Fix headers and footers
  4. Fix unnumbered headings (SLEP chapter, empty heading cleanup)
  5. Fix figure numbering gaps

Usage:
    python format_thesis.py                # dry run
    python format_thesis.py --apply        # apply changes
"""

import argparse
import copy
import datetime
import re
import shutil
import sys

try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

THESIS = "thesis_updated.docx"


def find_para(doc, needle, start=0):
    """Return (index, paragraph) for first paragraph containing needle."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if needle in p.text:
            return i, p
    return None, None


def replace_para_text(para, new_text):
    """Replace paragraph text preserving first run's formatting."""
    if not para.runs:
        para.text = new_text
        return
    for run in para.runs:
        run.text = ""
    para.runs[0].text = new_text


def insert_paragraph_after(doc, ref_para, text, style_name="Normal"):
    """Insert a new paragraph after ref_para with given style."""
    # Find the paragraph's XML element and insert after it
    ref_elem = ref_para._element
    new_para = docx.oxml.OxmlElement('w:p')

    # Set the style
    pPr = docx.oxml.OxmlElement('w:pPr')
    pStyle = docx.oxml.OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    new_para.append(pPr)

    # Add text run
    run = docx.oxml.OxmlElement('w:r')
    # Copy formatting from reference paragraph's first run if same style
    t = docx.oxml.OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    new_para.append(run)

    ref_elem.addnext(new_para)
    return new_para


def set_run_font(run_elem, font_name="Times New Roman", font_size=None, bold=None):
    """Set font properties on a run XML element."""
    rPr = run_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = docx.oxml.OxmlElement('w:rPr')
        run_elem.insert(0, rPr)

    if font_name:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = docx.oxml.OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)

    if font_size is not None:
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = docx.oxml.OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), str(font_size * 2))  # half-points

    if bold is not None:
        b = rPr.find(qn('w:b'))
        if bold:
            if b is None:
                b = docx.oxml.OxmlElement('w:b')
                rPr.append(b)
        else:
            if b is not None:
                rPr.remove(b)


# ══════════════════════════════════════════════════════════════════════
# STEP 1: Add Declaration and Acknowledgements
# ══════════════════════════════════════════════════════════════════════

def step1_front_matter(doc, dry_run=False):
    """Add Declaration and Acknowledgements pages after Abstract."""
    changes = 0

    # Find the keywords paragraph (last part of abstract section)
    kw_idx, kw_para = find_para(doc, "Keywords:")
    if kw_para is None:
        # Find end of abstract content
        kw_idx, kw_para = find_para(doc, "Adaptive Thresholding")
        if kw_para is None:
            print("  SKIP: Cannot find abstract end to insert after")
            return 0

    # Check if Declaration already exists
    decl_idx, _ = find_para(doc, "DECLARATION")
    if decl_idx is not None and decl_idx < 100:
        print("  SKIP: Declaration already exists")
    else:
        if not dry_run:
            # Insert acknowledgements first (will appear after declaration)
            ack_texts = [
                ("ACKNOWLEDGEMENTS", "Heading 1"),
                ("I would like to express my sincere gratitude to my supervisor for their guidance and support throughout this project. I also wish to thank my family and friends for their encouragement during this research.", "Normal"),
            ]
            # Insert declaration
            decl_texts = [
                ("DECLARATION", "Heading 1"),
                ("I declare that this is my own work and that where material is obtained from published or unpublished sources it has been fully acknowledged in the References section.", "Normal"),
                ("", "Normal"),
                ("Signature: ___________________________", "Normal"),
                ("Date: ___________________________", "Normal"),
                ("Student ID: w1912919", "Normal"),
            ]

            # Insert in reverse order (each after kw_para) so they appear in correct order
            all_inserts = list(reversed(decl_texts + ack_texts))
            ref = kw_para._element
            for text, style in all_inserts:
                new_elem = insert_paragraph_after(doc, kw_para, text, style)
                # The next insert should go after kw_para too (they stack)

        print("  ADDED: Declaration page (after abstract)")
        print("  ADDED: Acknowledgements page (after declaration)")
        changes += 2

    return changes


# ══════════════════════════════════════════════════════════════════════
# STEP 2: Restructure Abstract
# ══════════════════════════════════════════════════════════════════════

def step2_abstract(doc, dry_run=False):
    """Restructure abstract to template format: 3 paragraphs + descriptors + keywords."""
    changes = 0

    # Find the two abstract paragraphs
    idx1, p1 = find_para(doc, "Large Language Models (LLMs) are prone to hallucinations")
    idx2, p2 = find_para(doc, "A controlled three-condition experiment")

    if p1 is None:
        print("  SKIP: Abstract paragraph 1 not found")
        return 0

    # New abstract text — 3 paragraphs (Problem, Methodology, Results) ≤300 words
    problem = (
        "Problem: Large Language Models (LLMs) are prone to hallucinations — generating "
        "fluent but factually unsupported content. Existing verification frameworks apply "
        "uniform Natural Language Inference (NLI) thresholds regardless of the underlying "
        "evidence quality, leading to over-flagging of well-supported responses and "
        "under-flagging of poorly-supported ones."
    )

    methodology = (
        "Methodology: This project proposes Confidence-Weighted CONLI (Cw-CONLI), which "
        "dynamically adjusts the NLI verification threshold based on retrieval confidence. "
        "Three weighting functions (tiered, square-root, sigmoid) are implemented within "
        "AFLHR Lite, a two-layer RAG and verification pipeline. A v2 pipeline adds "
        "sliding-window NLI, sentence-level claim decomposition, and BGE embedding upgrades. "
        "Evaluation uses the HaluEval benchmark (20,000 samples across QA and Summarization "
        "tasks) in a three-condition experiment: C1 (RAG-only), C2 (static CONLI), and C3 "
        "(Cw-CONLI)."
    )

    results = (
        "Results: On the standard per-sample benchmark, C3 and C2 converge (combined F1: "
        "0.6988 vs 0.6998; McNemar's p = 1.0). Under realistic shared-index retrieval, "
        "the difference is significant (p = 1.4×10⁻⁵): C2 flags 100% of correct responses "
        "while C3 Sqrt reduces over-flagging to 44.9%. The primary contribution is v2 "
        "engineering — windowed NLI fixes summarisation (99% FPR to F1 ≈ 0.663), and claim "
        "decomposition catches partial hallucinations."
    )

    if not dry_run:
        replace_para_text(p1, problem)
        if p2:
            replace_para_text(p2, methodology)
            # Insert results paragraph after methodology
            insert_paragraph_after(doc, p2, results, "Normal")
        # Check if Subject Descriptors already exist
        sd_idx, _ = find_para(doc, "Subject Descriptors")
        if sd_idx is None:
            # Find keywords paragraph to insert before
            kw_idx, kw_para = find_para(doc, "Keywords:")
            if kw_para:
                # We need to insert Subject Descriptors before Keywords
                # Actually insert after the results paragraph
                pass
            # Add subject descriptors and keywords at end of abstract section
            # Find the last abstract paragraph
            last_abs = doc.paragraphs[idx2 + 1] if idx2 else p2
            sd_text = (
                "Subject Descriptors: Computing methodologies → Machine learning → "
                "Machine learning approaches; Information systems → Information retrieval "
                "→ Retrieval models and ranking"
            )
            insert_paragraph_after(doc, last_abs, sd_text, "Normal")

    print("  UPDATED: Abstract paragraph 1 (Problem)")
    print("  UPDATED: Abstract paragraph 2 (Methodology)")
    print("  ADDED: Abstract paragraph 3 (Results)")
    print("  ADDED: Subject Descriptors")
    changes += 4

    return changes


# ══════════════════════════════════════════════════════════════════════
# STEP 3: Fix Headers and Footers
# ══════════════════════════════════════════════════════════════════════

def step3_headers_footers(doc, dry_run=False):
    """Update headers and footers per template requirements."""
    changes = 0

    for i, section in enumerate(doc.sections):
        header = section.header
        footer = section.footer

        if not dry_run:
            # Clear and set header
            header.is_linked_to_previous = False
            for p in header.paragraphs:
                p.text = ""

            if header.paragraphs:
                hp = header.paragraphs[0]
                # Left-aligned: project name
                hp.text = ""
                run1 = hp.add_run("AFLHR Lite")
                run1.font.name = "Times New Roman"
                run1.font.size = Pt(10)

                # Add tab stop for right-aligned text
                run2 = hp.add_run("\tCw-CONLI Hallucination Detection")
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(10)

            # Clear and set footer
            footer.is_linked_to_previous = False
            for p in footer.paragraphs:
                p.text = ""

            if footer.paragraphs:
                fp = footer.paragraphs[0]
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = fp.add_run("Shaun Yogeshwaran — w1912919")
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

        print(f"  UPDATED: Section {i} header → 'AFLHR Lite | Cw-CONLI Hallucination Detection'")
        print(f"  UPDATED: Section {i} footer → 'Shaun Yogeshwaran — w1912919'")
        changes += 2

    return changes


# ══════════════════════════════════════════════════════════════════════
# STEP 4: Fix Unnumbered Headings
# ══════════════════════════════════════════════════════════════════════

def step4_headings(doc, dry_run=False):
    """Fix unnumbered headings and clean up empty heading paragraphs."""
    changes = 0

    # Specific fixes for SLEP chapter (Chapter 5)
    slep_fixes = {
        "Social Issues": "5.2.1 Social issues",
        "Legal Issues": "5.2.2 Legal issues",
        "Ethical Issues": "5.2.3 Ethical issues",
        "Professional Issues": "5.2.4 Professional issues",
        "Sustainability Considerations": "5.2.5 Sustainability considerations",
    }

    # Fix Algorithmic Pseudocode heading in Chapter 3
    algo_fixes = {
        "Algorithmic Pseudocode for Confidence-Weighted CONLI": "3.3.4 Algorithmic pseudocode for confidence-weighted CONLI",
    }

    # Fix Chapter 8 numbering issue (8.10 appears before 8.7)
    ch8_fixes = {
        "8.10 Latency Analysis Analysis": "8.6 Latency analysis",
    }

    all_fixes = {**slep_fixes, **algo_fixes, **ch8_fixes}

    for i, p in enumerate(doc.paragraphs):
        if not p.style.name.startswith("Heading"):
            continue

        text = p.text.strip()

        # Fix known unnumbered headings
        if text in all_fixes:
            new_text = all_fixes[text]
            if not dry_run:
                replace_para_text(p, new_text)
            print(f"  RENAMED: P{i} '{text}' → '{new_text}'")
            changes += 1

        # Clean up empty heading paragraphs (convert to Normal)
        elif text == "":
            old_style = p.style.name
            if not dry_run:
                p.style = doc.styles["Normal"]
            print(f"  CLEANED: P{i} empty {old_style} → Normal")
            changes += 1

    return changes


# ══════════════════════════════════════════════════════════════════════
# STEP 5: Fix Figure Numbering Gaps
# ══════════════════════════════════════════════════════════════════════

def step5_figures(doc, dry_run=False):
    """Fix figure numbering gaps — renumber sequentially within chapters."""
    changes = 0

    # First pass: collect all figure captions with their chapter
    figures = []
    current_chapter = 0
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1":
            m = re.match(r"CHAPTER\s+(\d+)", p.text, re.I)
            if m:
                current_chapter = int(m.group(1))

        m = re.match(r"Figure\s+(\d+)\.(\d+):", p.text)
        if m:
            old_ch = int(m.group(1))
            old_num = int(m.group(2))
            figures.append({
                "para_idx": i,
                "chapter": current_chapter,
                "old_label": f"Figure {old_ch}.{old_num}",
                "old_ch": old_ch,
                "old_num": old_num,
            })

    # Second pass: assign sequential numbers within each chapter
    chapter_counts = {}
    renames = {}
    for fig in figures:
        ch = fig["chapter"]
        chapter_counts[ch] = chapter_counts.get(ch, 0) + 1
        new_num = chapter_counts[ch]
        new_label = f"Figure {ch}.{new_num}"
        old_label = fig["old_label"]

        if old_label != new_label:
            renames[old_label] = new_label
            fig["new_label"] = new_label
        else:
            fig["new_label"] = None

    if not renames:
        print("  SKIP: All figures already numbered sequentially")
        return 0

    # Apply renames to captions and in-text references
    for old, new in renames.items():
        for i, p in enumerate(doc.paragraphs):
            if old in p.text:
                if not dry_run:
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                print(f"  RENAMED: P{i} '{old}' → '{new}'")
                changes += 1

    return changes


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Format thesis to template spec")
    parser.add_argument("--apply", action="store_true", help="Apply changes (default: dry run)")
    parser.add_argument("--no-backup", action="store_true", help="Skip backup")
    args = parser.parse_args()

    dry_run = not args.apply
    if dry_run:
        print("DRY RUN — no changes written. Use --apply to write.\n")

    doc = docx.Document(THESIS)

    if args.apply and not args.no_backup:
        backup = f"thesis_backup_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        shutil.copy2(THESIS, backup)
        print(f"Backup: {backup}\n")

    total = 0

    print("Step 1: Front Matter (Declaration + Acknowledgements)")
    total += step1_front_matter(doc, dry_run)

    print("\nStep 2: Abstract Restructuring")
    total += step2_abstract(doc, dry_run)

    print("\nStep 3: Headers and Footers")
    total += step3_headers_footers(doc, dry_run)

    print("\nStep 4: Heading Numbering")
    total += step4_headings(doc, dry_run)

    print("\nStep 5: Figure Numbering")
    print("  SKIPPED: Figure renumbering requires manual review.")
    print("  Known gaps: Figure 3.2 missing (3.1→3.3), Figure 7.11-7.12 missing (7.10→7.13)")
    print("  Note: Figures 7.x appear in Chapter 8 content — verify intended chapter numbering in Word.")

    if args.apply:
        doc.save(THESIS)
        print(f"\nSaved {THESIS} with {total} changes.")
        print("NOTE: Per-chapter headers (right side) require manual section breaks in Word.")
    else:
        print(f"\nDry run: {total} changes would be made. Use --apply to write.")


if __name__ == "__main__":
    main()
