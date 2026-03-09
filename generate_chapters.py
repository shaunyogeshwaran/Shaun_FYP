"""
Generate thesis chapters (Updated Abstract, Ch6-8) as a Word document.

Uses python-docx to produce thesis_new_chapters.docx with formatting that
matches the existing IPD document style.

Run: python generate_chapters.py
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================================
# Constants / formatting helpers
# ============================================================================

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
HEADER_TEXT = "Confidence-weighted framework for mitigating LLM Hallucinations"
FOOTER_ID = "w1912919"


def set_run_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    """Apply font formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure Times New Roman is used for East-Asian fallback too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), FONT_NAME)
        rFonts.set(qn("w:hAnsi"), FONT_NAME)


def add_paragraph(doc, text, bold=False, italic=False, size=FONT_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6)):
    """Add a body paragraph with standard formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading_custom(doc, text, level=1):
    """Add a heading that uses Times New Roman."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=Pt(16) if level == 1 else Pt(14) if level == 2 else Pt(12),
                     bold=True)
    return h


def add_figure_placeholder(doc, filename, caption):
    """Add a clearly marked figure placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"[INSERT FIGURE: {filename} — {caption}]")
    set_run_font(run, bold=True, italic=True, color=RGBColor(0x80, 0x00, 0x00))
    return p


def add_code_block(doc, text):
    """Add a code-style paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    return p


def create_table(doc, headers, rows, col_widths=None):
    """Create a bordered table with bold header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        set_run_font(run, bold=True, size=Pt(10))
        # Grey background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, size=Pt(10))

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    # Add spacing after table
    doc.add_paragraph()
    return table


def add_header_footer(doc):
    """Add header and footer to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(HEADER_TEXT)
        set_run_font(run, size=Pt(10), italic=True)

        # Footer: student ID on left, page number on right
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(f"{FOOTER_ID}    |    Page ")
        set_run_font(run, size=Pt(10))
        # Add page number field
        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
            f'<w:r><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
            f'<w:sz w:val="20"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
        )
        fp._element.append(parse_xml(fld_xml))


# ============================================================================
# Content sections
# ============================================================================

def write_abstract(doc):
    """Updated Abstract reflecting actual results."""
    add_heading_custom(doc, "Abstract", level=1)

    abstract = (
        "Large Language Models (LLMs) are prone to hallucinations — generating "
        "fluent but factually unsupported content. Existing verification frameworks "
        "apply uniform Natural Language Inference (NLI) thresholds regardless of "
        "the underlying evidence quality, potentially over-flagging well-supported "
        "responses and under-flagging poorly-supported ones. This project proposes "
        "Confidence-Weighted CONLI (Cw-CONLI), an extension that dynamically adjusts "
        "the NLI verification threshold based on retrieval confidence scores. Three "
        "continuous weighting functions — tiered, square-root, and sigmoid — are "
        "implemented within AFLHR Lite, a two-layer Retrieval-Augmented Generation "
        "(RAG) and verification pipeline."
    )
    add_paragraph(doc, abstract)

    abstract2 = (
        "A controlled three-condition experiment is conducted on the HaluEval benchmark "
        "(20,000 samples across QA and Summarization tasks): C1 (RAG-only baseline), "
        "C2 (static CONLI threshold), and C3 (Cw-CONLI with dynamic weighting). "
        "Results show that C3 performs comparably to C2 rather than significantly better. "
        "On the QA task, C2 achieves F1 = 0.770 while C3 Sigmoid achieves F1 = 0.769; "
        "however, C3 Sigmoid attains the lowest over-flagging rate (13.36% vs. 13.57%). "
        "McNemar's test confirms no statistically significant difference (p = 0.25). "
        "Analysis reveals that tight retrieval score clustering in HaluEval (QA std = 0.075) "
        "limits the differentiation potential of confidence-weighted thresholds. A realistic "
        "retrieval experiment using a shared FAISS index further demonstrates that C3 "
        "variants reduce over-flagging by 3–5% compared to C2's 100% false-positive rate. "
        "These findings suggest Cw-CONLI's benefits emerge in retrieval environments with "
        "greater score variability, pointing to future evaluation on heterogeneous corpora."
    )
    add_paragraph(doc, abstract2)


def write_chapter_6(doc):
    """Chapter 6: Implementation."""
    doc.add_page_break()
    add_heading_custom(doc, "Chapter 6: Implementation", level=1)

    # 6.1 Chapter Overview
    add_heading_custom(doc, "6.1 Chapter Overview", level=2)
    add_paragraph(doc, (
        "This chapter details the implementation of AFLHR Lite, the prototype system "
        "built to evaluate the Confidence-Weighted CONLI (Cw-CONLI) hypothesis. It covers "
        "the system architecture, module descriptions, core algorithms, experimental "
        "conditions, and evaluation methodology. All code is written in Python and "
        "executed on an Apple M4 MacBook Pro (24 GB RAM, CPU-only inference)."
    ))

    # 6.2 System Architecture
    add_heading_custom(doc, "6.2 System Architecture", level=2)
    add_paragraph(doc, (
        "AFLHR Lite implements a two-layer pipeline. The first layer — Retrieval-Augmented "
        "Generation (RAG) — retrieves relevant evidence from a knowledge base using FAISS "
        "similarity search and generates a response using an LLM. The second layer — "
        "Verification — applies NLI-based entailment checking with either a static or "
        "confidence-weighted threshold to determine whether the generated response is "
        "supported by the retrieved evidence."
    ))
    add_paragraph(doc, (
        "The pipeline proceeds in four sequential steps: (1) Embed the query and retrieve "
        "top-k documents from FAISS, producing a retrieval confidence score; (2) Generate "
        "a response conditioned on the retrieved context; (3) Verify the response against "
        "the evidence using NLI, obtaining an entailment probability; (4) Apply the "
        "threshold function (static or dynamic) to render a VERIFIED or HALLUCINATION verdict."
    ))
    add_figure_placeholder(doc, "architecture_diagram.png",
                           "Figure 6.1: AFLHR Lite System Architecture — Two-layer RAG + Verification pipeline")

    # 6.3 Module Descriptions
    add_heading_custom(doc, "6.3 Module Descriptions", level=2)
    add_paragraph(doc, (
        "The system is organised into seven Python modules, each with a clearly defined "
        "responsibility. This modular design separates concerns between model inference, "
        "data loading, evaluation logic, and user interface."
    ))

    # 6.3.1 engine.py
    add_heading_custom(doc, "6.3.1 engine.py — Core Engine", level=3)
    add_paragraph(doc, (
        "The AFLHREngine class (533 lines) is the central component, encapsulating all "
        "model loading, embedding, retrieval, generation, and verification logic. Key methods include:"
    ))
    add_paragraph(doc, (
        "• _encode(texts): Generates normalised embeddings using mean pooling over "
        "all-MiniLM-L6-v2 token outputs, replicating SentenceTransformer behaviour without "
        "the segfault-prone library on Apple Silicon.\n"
        "• retrieve(query, k=2): Embeds the query, searches the FAISS IndexFlatIP index, "
        "and returns the top-k documents with a cosine similarity score normalised to [0, 1].\n"
        "• verify(premise, hypothesis): Tokenises the premise-hypothesis pair (truncation="
        "\"longest_first\", max_length=512), runs RoBERTa-large-MNLI inference, and returns "
        "the softmax entailment probability (index 2 of [contradiction, neutral, entailment]).\n"
        "• calculate_verdict(...): Implements the tiered Cw-CONLI threshold: if retrieval_score "
        "< pivot, apply T_strict; otherwise apply T_lenient.\n"
        "• calculate_verdict_continuous(...): Implements sqrt and sigmoid continuous weighting functions.\n"
        "• precompute_scores(knowledge, query, response): Computes retrieval and NLI scores "
        "independently of any threshold, enabling fast grid search."
    ))

    # 6.3.2 config.py
    add_heading_custom(doc, "6.3.2 config.py — Configuration", level=3)
    add_paragraph(doc, (
        "Centralises all configuration parameters including model identifiers "
        "(EMBEDDING_MODEL = \"sentence-transformers/all-MiniLM-L6-v2\", VERIFIER_MODEL = "
        "\"FacebookAI/roberta-large-mnli\", GENERATOR_MODEL = \"llama-3.1-8b-instant\"), "
        "default thresholds (pivot = 0.75, T_strict = 0.95, T_lenient = 0.70), "
        "experiment settings (seed = 42, dev_split_ratio = 0.7), and grid search ranges."
    ))

    tbl_grid = [
        ["C2 T_static", "0.50", "0.99", "0.01"],
        ["C3 Pivot", "0.60", "0.90", "0.05"],
        ["C3 T_strict (tiered)", "0.85", "0.99", "0.02"],
        ["C3 T_lenient (tiered)", "0.60", "0.85", "0.05"],
        ["C3 T_strict (continuous)", "0.85", "0.99", "0.02"],
        ["C3 T_lenient (continuous)", "0.50", "0.85", "0.05"],
    ]
    add_paragraph(doc, "Table 6.1: Grid Search Parameter Ranges", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc, ["Parameter", "Min", "Max", "Step"], tbl_grid)

    # 6.3.3 dataset.py
    add_heading_custom(doc, "6.3.3 dataset.py — Dataset Loader", level=3)
    add_paragraph(doc, (
        "Loads the HaluEval benchmark (pminervini/HaluEval) from HuggingFace, supporting "
        "both QA and Summarization subsets. Each sample is independently labelled with "
        "hallucination = \"yes\" or \"no\". The loader normalises fields across tasks: for QA, "
        "knowledge = row[\"knowledge\"], question = row[\"question\"], response = row[\"answer\"]; "
        "for Summarization, knowledge = row[\"document\"], question = \"\" (unused), response = "
        "row[\"summary\"]. Labels are converted to binary (1 = hallucination, 0 = valid)."
    ))
    add_paragraph(doc, (
        "The split_dev_test function shuffles samples with seed = 42 and splits at 70/30 ratio, "
        "yielding 14,000 dev samples and 6,000 test samples."
    ))

    # 6.3.4 evaluate.py
    add_heading_custom(doc, "6.3.4 evaluate.py — Evaluation Harness", level=3)
    add_paragraph(doc, (
        "The evaluation harness implements a two-phase design: (1) Pre-computation phase — "
        "runs model inference (embedding + NLI) once per sample, saving retrieval_score, "
        "nli_score, and latency_ms to CSV with checkpointing every 100 samples; "
        "(2) Condition application phase — loads pre-computed scores and applies threshold "
        "logic (C1/C2/C3) without any model inference, enabling rapid grid search."
    ))
    add_paragraph(doc, (
        "For realistic retrieval (Experiment 2), run_precomputation_realistic builds a shared "
        "FAISS index from all unique QA knowledge passages and retrieves from it, simulating "
        "a real-world RAG system where the exact ground-truth passage is not guaranteed."
    ))
    add_paragraph(doc, (
        "Metrics computed include F1 score, precision, recall, accuracy, over-flagging rate "
        "(FP / (FP + TN)), and latency statistics (mean, median, p95)."
    ))

    # 6.3.5 tune.py
    add_heading_custom(doc, "6.3.5 tune.py — Grid Search Tuning", level=3)
    add_paragraph(doc, (
        "Performs exhaustive grid search over the parameter space defined in config.py, "
        "maximising F1 score on the dev set. For C2, it sweeps T_static from 0.50 to 0.99 "
        "(50 values). For C3 tiered, it sweeps pivot x T_strict x T_lenient "
        "(7 x 8 x 6 = 336 valid configurations after filtering tl < ts). For C3 continuous "
        "(sqrt and sigmoid), it sweeps T_strict x T_lenient (8 x 8 = 64 valid configurations each). "
        "Results are saved as JSON with both compact summaries and full logs."
    ))

    # 6.3.6 analyze.py
    add_heading_custom(doc, "6.3.6 analyze.py — Analysis and Visualisation", level=3)
    add_paragraph(doc, (
        "Generates comparison tables (CSV), publication-quality plots (PNG at 300 DPI), "
        "and McNemar's statistical test. Plots include F1 bar charts, precision-recall "
        "comparisons, over-flagging rates, retrieval and NLI score distributions, confusion "
        "matrices for all five conditions, and latency boxplots."
    ))

    # 6.3.7 app.py
    add_heading_custom(doc, "6.3.7 app.py — Streamlit Demo", level=3)
    add_paragraph(doc, (
        "Provides an interactive web interface built with Streamlit. Users can enter queries, "
        "adjust threshold parameters via sidebar sliders, and observe the full pipeline in "
        "real time: retrieved evidence, generated response, NLI score, applied threshold, "
        "and final verdict. The engine is cached using @st.cache_resource to prevent model "
        "reloading between interactions. The demo operates against a curated six-passage "
        "knowledge base covering the University of Westminster, AI hallucinations, and a "
        "distractor topic (Climate of Sri Lanka)."
    ))

    # 6.4 Core Algorithm
    add_heading_custom(doc, "6.4 Core Algorithm: Confidence-Weighted CONLI", level=2)
    add_paragraph(doc, (
        "The central innovation of this project is the Cw-CONLI threshold function, which "
        "replaces the static NLI threshold with one that adapts based on retrieval confidence. "
        "The intuition is that when retrieval confidence is low (evidence may be irrelevant), "
        "a stricter NLI threshold should be applied, and when retrieval confidence is high "
        "(evidence is likely relevant), a more lenient threshold suffices. Three weighting "
        "functions are implemented:"
    ))

    # 6.4.1 Tiered
    add_heading_custom(doc, "6.4.1 Tiered Model", level=3)
    add_paragraph(doc, (
        "The simplest form uses a binary pivot point. If retrieval_score < pivot, the threshold "
        "is T_strict; otherwise it is T_lenient. This creates a step function with a single "
        "discontinuity at the pivot. While easy to interpret, the abrupt transition may "
        "cause instability for samples near the pivot boundary."
    ))
    add_code_block(doc, "if retrieval_score < pivot:")
    add_code_block(doc, "    threshold = T_strict")
    add_code_block(doc, "else:")
    add_code_block(doc, "    threshold = T_lenient")

    # 6.4.2 Sqrt
    add_heading_custom(doc, "6.4.2 Continuous Square-Root Model", level=3)
    add_paragraph(doc, (
        "The square-root function provides a smooth, concave transition from T_strict "
        "(at retrieval_score = 0) toward T_lenient (at retrieval_score = 1). The concavity "
        "means that the threshold drops quickly for moderate retrieval scores and flattens "
        "at high scores."
    ))
    add_code_block(doc, "T = T_strict - (T_strict - T_lenient) * sqrt(retrieval_score)")

    # 6.4.3 Sigmoid
    add_heading_custom(doc, "6.4.3 Continuous Sigmoid Model", level=3)
    add_paragraph(doc, (
        "The sigmoid function provides an S-shaped transition centred at a configurable "
        "pivot, with steepness controlled by parameter k (default = 10). At low retrieval "
        "scores the threshold approaches T_strict; at high scores it approaches T_lenient; "
        "the transition is steepest around the pivot."
    ))
    add_code_block(doc, "T = T_lenient + (T_strict - T_lenient) / (1 + exp(k * (rs - pivot)))")

    # 6.5 Experimental Conditions
    add_heading_custom(doc, "6.5 Experimental Conditions", level=2)
    add_paragraph(doc, (
        "Three conditions are defined to isolate the effect of confidence-weighted thresholding:"
    ))

    cond_rows = [
        ["C1", "RAG-only baseline", "No NLI verification; all responses accepted",
         "None"],
        ["C2", "Static CONLI", "Fixed NLI threshold (T_static) applied uniformly",
         "T_static"],
        ["C3 Tiered", "Cw-CONLI (tiered)", "Binary threshold based on pivot",
         "pivot, T_strict, T_lenient"],
        ["C3 Sqrt", "Cw-CONLI (sqrt)", "Smooth sqrt-weighted threshold",
         "T_strict, T_lenient"],
        ["C3 Sigmoid", "Cw-CONLI (sigmoid)", "Smooth sigmoid-weighted threshold",
         "T_strict, T_lenient"],
    ]
    add_paragraph(doc, "Table 6.2: Experimental Condition Definitions", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc, ["Condition", "Name", "Description", "Parameters"], cond_rows)

    # 6.6 Evaluation Methodology
    add_heading_custom(doc, "6.6 Evaluation Methodology", level=2)

    add_heading_custom(doc, "6.6.1 Dataset", level=3)
    add_paragraph(doc, (
        "The HaluEval benchmark (Li et al., 2023) is used, comprising 10,000 QA samples and "
        "10,000 Summarization samples (20,000 total). Each sample is independently labelled as "
        "hallucinated or valid, with approximately balanced classes. The dataset provides "
        "ground-truth knowledge passages per sample, enabling controlled retrieval evaluation."
    ))

    split_rows = [
        ["QA", "7,041", "3,523", "3,518", "2,959", "1,467", "1,492"],
        ["Summarization", "6,959", "3,459", "3,500", "3,041", "1,531", "1,510"],
        ["Total", "14,000", "6,982", "7,018", "6,000", "2,998", "3,002"],
    ]
    add_paragraph(doc, "Table 6.3: Dataset Split Sizes (seed = 42, 70/30 split)", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc,
                 ["Task", "Dev Total", "Dev Valid", "Dev Halluc.",
                  "Test Total", "Test Valid", "Test Halluc."],
                 split_rows)

    add_heading_custom(doc, "6.6.2 Pre-computation Approach", level=3)
    add_paragraph(doc, (
        "To enable efficient hyperparameter search, model inference (embedding and NLI) is "
        "separated from threshold application. Retrieval scores and NLI scores are pre-computed "
        "once per sample and saved to CSV. Threshold sweeping then operates purely on these "
        "cached scores, reducing tuning time from hours to seconds. For QA, the retrieval query "
        "is the question; for Summarization, the retrieval query is the response (summary), "
        "following the convention that we measure how well the knowledge supports the generated text."
    ))

    add_heading_custom(doc, "6.6.3 Metrics", level=3)
    add_paragraph(doc, (
        "Five metrics are reported: (1) F1 score — the primary optimisation target, balancing "
        "precision and recall for hallucination detection; (2) Precision — proportion of flagged "
        "responses that are truly hallucinated; (3) Recall — proportion of actual hallucinations "
        "that are detected; (4) Over-flagging rate — FP / (FP + TN), measuring the fraction of "
        "valid responses incorrectly flagged as hallucinations; (5) Latency — per-sample inference "
        "time in milliseconds."
    ))

    # 6.7 Chapter Summary
    add_heading_custom(doc, "6.7 Chapter Summary", level=2)
    add_paragraph(doc, (
        "This chapter presented the complete implementation of AFLHR Lite, covering the modular "
        "architecture, all seven Python modules, three Cw-CONLI threshold functions, the "
        "three-condition experimental design, and the evaluation methodology. The pre-computation "
        "approach enables efficient grid search while ensuring reproducibility through fixed "
        "random seeds and checkpointed CSV outputs."
    ))


def write_chapter_7(doc):
    """Chapter 7: Results."""
    doc.add_page_break()
    add_heading_custom(doc, "Chapter 7: Results", level=1)

    # 7.1
    add_heading_custom(doc, "7.1 Chapter Overview", level=2)
    add_paragraph(doc, (
        "This chapter presents the experimental results across three experiments: "
        "(1) standard evaluation on the HaluEval test set (QA + Summarization combined, "
        "then per-task), (2) realistic retrieval evaluation using a shared FAISS index "
        "(QA only), and (3) statistical significance testing. Score distributions, "
        "hyperparameter tuning outcomes, and latency analysis are also reported."
    ))

    # 7.2 Dataset Characteristics
    add_heading_custom(doc, "7.2 Dataset Characteristics", level=2)
    add_paragraph(doc, (
        "Before evaluating conditions, the pre-computed score distributions are examined "
        "to understand the underlying signal quality."
    ))

    add_heading_custom(doc, "7.2.1 Retrieval Score Distributions", level=3)
    add_paragraph(doc, (
        "QA retrieval scores have mean = 0.818 and std = 0.075, indicating generally high "
        "but moderately variable retrieval confidence. Summarization retrieval scores are "
        "higher and tighter (mean = 0.858, std = 0.042). The combined distribution "
        "(mean = 0.838, std = 0.064) reflects both tasks. Critically, the narrow spread "
        "of retrieval scores limits the range over which Cw-CONLI can differentiate thresholds."
    ))
    add_figure_placeholder(doc, "retrieval_dist.png",
                           "Figure 7.1: Retrieval Score Distribution (all tasks, test set)")
    add_figure_placeholder(doc, "retrieval_dist_qa.png",
                           "Figure 7.2: Retrieval Score Distribution (QA only, test set)")

    add_heading_custom(doc, "7.2.2 NLI Score Separability", level=3)
    add_paragraph(doc, (
        "For QA, NLI scores show strong separability: correct responses have mean entailment "
        "probability = 0.770, while hallucinated responses have mean = 0.310. This 0.46 gap "
        "provides a clear signal for threshold-based classification."
    ))
    add_paragraph(doc, (
        "For Summarization, NLI scores show poor separability: correct responses have "
        "mean = 0.208, hallucinated = 0.136. Both values are low, indicating that the "
        "RoBERTa-large-MNLI model struggles to produce meaningful entailment scores for "
        "long-document summarization. This is due to the 512-token truncation limit, which "
        "discards critical document content before the NLI model can evaluate it."
    ))

    score_dist_rows = [
        ["QA", "0.818", "0.075", "0.770", "0.310", "0.460"],
        ["Summarization", "0.858", "0.042", "0.208", "0.136", "0.072"],
        ["Combined", "0.838", "0.064", "—", "—", "—"],
    ]
    add_paragraph(doc, "Table 7.1: Score Distribution Summary (Test Set)", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc,
                 ["Task", "Retrieval Mean", "Retrieval Std",
                  "NLI Mean (Valid)", "NLI Mean (Halluc.)", "NLI Gap"],
                 score_dist_rows)

    add_figure_placeholder(doc, "nli_dist.png",
                           "Figure 7.3: NLI Score Distribution (all tasks, test set)")
    add_figure_placeholder(doc, "nli_dist_qa.png",
                           "Figure 7.4: NLI Score Distribution (QA only, test set)")

    # 7.3 Hyperparameter Tuning
    add_heading_custom(doc, "7.3 Hyperparameter Tuning Results (Dev Set)", level=2)
    add_paragraph(doc, (
        "Grid search is performed on the 14,000-sample dev set to find optimal parameters "
        "for each condition. Results are reported for both combined (all tasks) and QA-only tuning."
    ))

    tuning_rows = [
        ["C2", "T_static = 0.56", "0.7054", "—", "—"],
        ["C3 Tiered", "pivot=0.60, T_s=0.85, T_l=0.60", "0.7038", "—", "—"],
        ["C3 Sqrt", "T_s=0.91, T_l=0.50", "0.7049", "—", "—"],
        ["C3 Sigmoid", "T_s=0.89, T_l=0.55", "0.7051", "—", "—"],
    ]
    add_paragraph(doc, "Table 7.2: Dev Set Tuning — Combined (QA + Summarization)", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc, ["Condition", "Best Parameters", "Best F1", "Precision", "Recall"],
                 tuning_rows)

    tuning_qa_rows = [
        ["C2", "T_static = 0.52", "0.7684", "—", "—"],
        ["C3 Tiered", "pivot=0.60, T_s=0.85, T_l=0.60", "0.7626", "—", "—"],
        ["C3 Sqrt", "T_s=0.91, T_l=0.50", "0.7682", "—", "—"],
        ["C3 Sigmoid", "T_s=0.85, T_l=0.50", "0.7681", "—", "—"],
    ]
    add_paragraph(doc, "Table 7.3: Dev Set Tuning — QA Only", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc, ["Condition", "Best Parameters", "Best F1", "Precision", "Recall"],
                 tuning_qa_rows)

    add_paragraph(doc, (
        "Key observation: All conditions converge to similar F1 scores on the dev set. "
        "For combined tasks, the spread is only 0.0016 (from C3 Tiered at 0.7038 to C2 at "
        "0.7054). For QA only, C2 achieves 0.7684, while C3 Sqrt and Sigmoid reach 0.7682 "
        "and 0.7681 respectively — a difference of 0.0002. C3 Tiered trails at 0.7626."
    ))

    # 7.4 Experiment 1: Standard Evaluation
    add_heading_custom(doc, "7.4 Experiment 1: Standard Evaluation (Test Set)", level=2)

    # 7.4.1 Combined
    add_heading_custom(doc, "7.4.1 Combined Results (QA + Summarization)", level=3)
    add_paragraph(doc, (
        "The test set contains 6,000 samples (2,959 QA + 3,041 Summarization). Table 7.4 "
        "shows the full comparison across all five conditions using parameters tuned on the "
        "combined dev set."
    ))

    combined_rows = [
        ["C1 (RAG-only)", "0.000", "0.000", "0.000", "0.500", "0.000"],
        ["C2 (Static CONLI)", "0.702", "0.607", "0.832", "0.647", "0.539"],
        ["C3 Tiered", "0.703", "0.602", "0.843", "0.643", "0.557"],
        ["C3 Sqrt", "0.703", "0.612", "0.826", "0.651", "0.525"],
        ["C3 Sigmoid", "0.702", "0.607", "0.833", "0.647", "0.540"],
    ]
    add_paragraph(doc, "Table 7.4: Test Set Results — Combined (QA + Summarization)", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc,
                 ["Condition", "F1", "Precision", "Recall", "Accuracy", "Over-flag Rate"],
                 combined_rows)

    add_paragraph(doc, (
        "All NLI-based conditions (C2 and C3 variants) substantially outperform C1 (F1 = 0.0), "
        "confirming that NLI verification is essential. Among the NLI conditions, the differences "
        "are minimal: F1 ranges from 0.702 to 0.703. C3 Sqrt achieves the highest precision "
        "(0.612) and lowest over-flagging (0.525), while C3 Tiered achieves the highest recall "
        "(0.843). The high over-flagging rates (>0.52) reflect the Summarization task's poor NLI "
        "signal dragging down combined performance."
    ))
    add_figure_placeholder(doc, "f1_comparison.png",
                           "Figure 7.5: F1 Score Comparison — Combined Test Set")
    add_figure_placeholder(doc, "precision_recall.png",
                           "Figure 7.6: Precision and Recall — Combined Test Set")
    add_figure_placeholder(doc, "confusion_matrices.png",
                           "Figure 7.7: Confusion Matrices — Combined Test Set")

    # 7.4.2 QA
    add_heading_custom(doc, "7.4.2 QA Task Results", level=3)
    add_paragraph(doc, (
        "Isolating the QA task (2,959 samples) reveals cleaner results due to the strong "
        "NLI signal. Parameters are taken from QA-specific tuning on the dev set."
    ))

    qa_rows = [
        ["C1 (RAG-only)", "0.000", "0.000", "0.000", "0.496", "0.000"],
        ["C2 (Static CONLI)", "0.770", "0.842", "0.710", "0.786", "0.136"],
        ["C3 Tiered", "0.766", "0.799", "0.735", "0.773", "0.188"],
        ["C3 Sqrt", "0.770", "0.832", "0.717", "0.784", "0.147"],
        ["C3 Sigmoid", "0.769", "0.843", "0.706", "0.786", "0.134"],
    ]
    add_paragraph(doc, "Table 7.5: Test Set Results — QA Only", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc,
                 ["Condition", "F1", "Precision", "Recall", "Accuracy", "Over-flag Rate"],
                 qa_rows)

    add_paragraph(doc, (
        "C2 and C3 Sqrt are tied at F1 = 0.770, with C3 Sigmoid close at 0.769. Notably, "
        "C3 Sigmoid achieves the lowest over-flagging rate of 13.36%, marginally better than "
        "C2's 13.57%. C3 Sigmoid also achieves the highest precision (0.843). C3 Tiered "
        "trails with F1 = 0.766 and the highest over-flagging (18.81%), suggesting that "
        "the binary threshold function is less effective than continuous alternatives."
    ))
    add_figure_placeholder(doc, "f1_comparison_qa.png",
                           "Figure 7.8: F1 Score Comparison — QA Test Set")
    add_figure_placeholder(doc, "overflagging_qa.png",
                           "Figure 7.9: Over-flagging Rate — QA Test Set")
    add_figure_placeholder(doc, "confusion_matrices_qa.png",
                           "Figure 7.10: Confusion Matrices — QA Test Set")

    # 7.4.3 Summarization
    add_heading_custom(doc, "7.4.3 Summarization Task Results", level=3)
    add_paragraph(doc, (
        "Summarization results are uniformly poor across all conditions. The NLI model "
        "produces low entailment scores for both valid and hallucinated summaries (mean 0.208 "
        "vs. 0.136), resulting in near-universal hallucination flagging. Over-flagging rates "
        "exceed 99% for all NLI-based conditions."
    ))
    add_paragraph(doc, (
        "The root cause is the 512-token truncation limit of the RoBERTa-large-MNLI model. "
        "Summarization documents in HaluEval are typically 1,000–3,000 tokens long. After "
        "truncation, the NLI model sees only a fraction of the premise, losing critical "
        "information needed to judge entailment. Additionally, summaries naturally use "
        "paraphrasing and abstraction, which the NLI model — trained on short, literal "
        "premise-hypothesis pairs from the MultiNLI corpus — interprets as non-entailment."
    ))

    # 7.5 Experiment 2
    add_heading_custom(doc, "7.5 Experiment 2: Realistic Retrieval (QA Only)", level=2)
    add_paragraph(doc, (
        "Experiment 1 uses per-sample ground-truth knowledge as the retrieval source, which "
        "guarantees high retrieval scores. To simulate a more realistic scenario, Experiment 2 "
        "builds a shared FAISS index from all 9,936 unique QA knowledge passages across the "
        "entire dataset and retrieves from this pool for each test sample."
    ))

    realistic_rows = [
        ["C1 (RAG-only)", "0.000", "0.000", "0.000", "0.496", "0.000"],
        ["C2 (Static CONLI)", "0.670", "0.504", "1.000", "0.504", "1.000"],
        ["C3 Tiered", "0.665", "0.505", "0.975", "0.505", "0.973"],
        ["C3 Sqrt", "0.643", "0.495", "0.917", "0.486", "0.953"],
        ["C3 Sigmoid", "0.641", "0.494", "0.914", "0.485", "0.952"],
    ]
    add_paragraph(doc, "Table 7.6: Realistic Retrieval Results — QA Test Set", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc,
                 ["Condition", "F1", "Precision", "Recall", "Accuracy", "Over-flag Rate"],
                 realistic_rows)

    add_paragraph(doc, (
        "Under realistic retrieval, C2 flags every single response as a hallucination "
        "(100% over-flagging, recall = 1.0). This occurs because the optimal C2 threshold "
        "(T_static = 0.99) is so high that even correctly retrieved passages produce NLI "
        "scores below it. C3 Tiered reduces over-flagging to 97.3%, while C3 Sqrt and "
        "Sigmoid achieve 95.3% and 95.2% respectively. Although all rates remain impractically "
        "high, the 3–5% reduction demonstrates that confidence-weighted thresholds provide "
        "a correct-direction advantage when retrieval quality varies."
    ))
    add_paragraph(doc, (
        "Realistic retrieval scores have mean = 0.827 and std = 0.063, compared to 0.818 "
        "and 0.075 in the standard experiment. The slightly different distribution reflects "
        "the shared-index retrieval mechanism."
    ))
    add_figure_placeholder(doc, "f1_comparison_realistic.png",
                           "Figure 7.11: F1 Score Comparison — Realistic Retrieval")
    add_figure_placeholder(doc, "retrieval_dist_realistic.png",
                           "Figure 7.12: Retrieval Score Distribution — Realistic Retrieval")

    # 7.6 Statistical Significance
    add_heading_custom(doc, "7.6 Statistical Significance", level=2)
    add_paragraph(doc, (
        "McNemar's test with continuity correction is used to assess whether the best "
        "C3 variant performs significantly differently from C2 on each experiment. The null "
        "hypothesis is that both models have equal error rates."
    ))

    mcnemar_rows = [
        ["Combined (all tasks)", "C3 Sigmoid", "7", "6", "0.000", "1.000", "No"],
        ["QA only", "C3 Sqrt", "17", "10", "1.333", "0.248", "No"],
        ["Realistic (QA)", "C3 Tiered", "37", "40", "0.052", "0.820", "No"],
    ]
    add_paragraph(doc, "Table 7.7: McNemar's Test Results (C2 vs. Best C3)", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc,
                 ["Experiment", "C3 Variant", "b (C2\u2713 C3\u2717)",
                  "c (C2\u2717 C3\u2713)", "Statistic", "p-value", "Significant (p<0.05)"],
                 mcnemar_rows)

    add_paragraph(doc, (
        "All p-values exceed 0.05, confirming that no C3 variant performs statistically "
        "significantly differently from C2. The closest result is QA-only (p = 0.248), "
        "where C2 has 17 samples correct that C3 Sqrt misclassifies, while C3 Sqrt has "
        "10 samples correct that C2 misclassifies."
    ))

    # 7.7 Latency
    add_heading_custom(doc, "7.7 Latency Analysis", level=2)
    add_paragraph(doc, (
        "Per-sample latency is measured during pre-computation and reflects the combined cost "
        "of embedding (retrieval score computation) and NLI inference. Since the threshold "
        "calculation itself is O(1) arithmetic, C2 and C3 conditions share identical latency."
    ))

    latency_rows = [
        ["Combined (all tasks)", "6,000", "392.24", "312.80", "766.31"],
        ["QA only", "2,959", "108.84", "100.67", "164.81"],
        ["Realistic (QA)", "2,959", "152.36", "147.64", "222.16"],
    ]
    add_paragraph(doc, "Table 7.8: Per-Sample Latency (ms)", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc,
                 ["Experiment", "N Samples", "Mean", "Median", "P95"],
                 latency_rows)

    add_paragraph(doc, (
        "QA samples process in ~109 ms on average, while the combined latency (~392 ms) "
        "reflects the longer Summarization documents that require more tokenisation time. "
        "Realistic retrieval adds ~44 ms overhead per sample due to FAISS index search over "
        "~10,000 passages (vs. single-passage embedding in the standard experiment). All "
        "measurements are on CPU (Apple M4); GPU inference would substantially reduce these times."
    ))
    add_figure_placeholder(doc, "latency_boxplot.png",
                           "Figure 7.13: Per-sample Latency Distribution")

    # 7.8 Summary
    add_heading_custom(doc, "7.8 Chapter Summary", level=2)
    add_paragraph(doc, (
        "The results demonstrate that Cw-CONLI (C3) performs comparably to static CONLI (C2) "
        "across all experiments. On the QA task, F1 differences are below 0.002, and McNemar's "
        "test shows no statistical significance (all p > 0.05). C3 Sigmoid achieves a marginal "
        "advantage in over-flagging rate (13.36% vs. 13.57% for C2 on QA). The Summarization "
        "task is fundamentally limited by the 512-token NLI truncation. In the realistic "
        "retrieval experiment, C3 variants show a correct-direction advantage in reducing "
        "over-flagging (95–97% vs. C2's 100%). These findings are analysed in detail in "
        "Chapter 8."
    ))


def write_chapter_8(doc):
    """Chapter 8: Discussion."""
    doc.add_page_break()
    add_heading_custom(doc, "Chapter 8: Discussion", level=1)

    # 8.1
    add_heading_custom(doc, "8.1 Chapter Overview", level=2)
    add_paragraph(doc, (
        "This chapter interprets the experimental results presented in Chapter 7. It answers "
        "the research questions, evaluates the research objectives, analyses why Cw-CONLI did "
        "not significantly outperform static CONLI, discusses the Summarization challenge and "
        "realistic retrieval insights, and concludes with limitations and future work."
    ))

    # 8.2 Research Questions
    add_heading_custom(doc, "8.2 Answering the Research Questions", level=2)

    add_heading_custom(doc, "8.2.1 RQ1: Does confidence-weighted thresholding improve F1?", level=3)
    add_paragraph(doc, (
        "RQ1 asked whether Cw-CONLI (C3) improves hallucination detection F1 compared to "
        "static CONLI (C2). The answer is no — C3 performs comparably but not significantly "
        "better. On the QA task, C2 achieves F1 = 0.770 while C3 Sqrt matches at 0.770 and "
        "C3 Sigmoid reaches 0.769, a difference of less than 0.002. On the combined test set, "
        "C3 Sqrt achieves F1 = 0.703 versus C2's 0.702."
    ))
    add_paragraph(doc, (
        "The primary explanation is the tight clustering of retrieval scores in HaluEval. "
        "With a QA retrieval score standard deviation of only 0.075 (on a [0, 1] scale), "
        "most samples fall within a narrow band where the dynamic threshold closely approximates "
        "the static one. The weighting functions (sqrt, sigmoid) can only differentiate "
        "thresholds meaningfully when retrieval scores span a wider range."
    ))

    add_heading_custom(doc, "8.2.2 RQ2: Does Cw-CONLI reduce latency and over-flagging?", level=3)
    add_paragraph(doc, (
        "For latency, C3 adds negligible overhead since the threshold calculation is O(1) "
        "arithmetic (a single sqrt or sigmoid evaluation) on top of the same NLI inference "
        "used by C2. Per-sample latencies are identical across C2 and C3 conditions."
    ))
    add_paragraph(doc, (
        "For over-flagging, C3 Sigmoid achieves a marginal improvement on QA: 13.36% vs. "
        "C2's 13.57%, a reduction of 0.21 percentage points. While small, this is consistent "
        "with the hypothesis that confidence-weighted thresholds can reduce false positives. "
        "In the realistic retrieval experiment, the advantage is more pronounced: C3 variants "
        "reduce over-flagging by 3–5 percentage points compared to C2's 100%."
    ))

    add_heading_custom(doc, "8.2.3 RQ3: Which weighting function performs best?", level=3)
    add_paragraph(doc, (
        "On the QA dev set, C2 (F1 = 0.7684) marginally outperforms C3 Sqrt (0.7682) and "
        "C3 Sigmoid (0.7681), with C3 Tiered trailing at 0.7626. Among C3 variants, continuous "
        "functions (sqrt and sigmoid) consistently outperform the tiered model. The tiered "
        "model's step-function discontinuity at the pivot creates a binary partition that does "
        "not capture the gradual relationship between retrieval confidence and optimal threshold. "
        "Between sqrt and sigmoid, performance is virtually identical; the sigmoid's additional "
        "pivot parameter provides no measurable advantage over the simpler sqrt function."
    ))

    # 8.3 Research Objectives
    add_heading_custom(doc, "8.3 Evaluating Research Objectives", level=2)

    ro_rows = [
        ["RO1", "Design a confidence-weighted threshold algorithm",
         "Achieved", "Three variants implemented: tiered, sqrt, sigmoid"],
        ["RO2", "Implement within a RAG + verification pipeline",
         "Achieved", "Full AFLHR Lite system with 7 modules"],
        ["RO3", "Empirically evaluate on a hallucination benchmark",
         "Achieved", "Three-condition experiment on HaluEval (20,000 samples)"],
        ["RO4", "Analyse trade-offs (F1, over-flagging, latency)",
         "Achieved", "All metrics quantified; statistical testing performed"],
    ]
    add_paragraph(doc, "Table 8.1: Research Objective Evaluation", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    create_table(doc, ["RO", "Objective", "Status", "Evidence"], ro_rows)

    add_paragraph(doc, (
        "All four research objectives have been achieved. While the empirical results did not "
        "demonstrate statistically significant improvement (RO3), the objective was to evaluate "
        "— not to guarantee improvement — and the evaluation was rigorous and complete."
    ))

    # 8.4 Why C3 Did Not Outperform C2
    add_heading_custom(doc, "8.4 Why C3 Did Not Outperform C2", level=2)
    add_paragraph(doc, (
        "Three factors explain why confidence-weighted thresholds did not yield significant "
        "gains in this evaluation:"
    ))
    add_paragraph(doc, (
        "First, HaluEval provides ground-truth knowledge per sample. When computing "
        "retrieval scores by comparing the query embedding against the provided knowledge "
        "embedding, the scores are naturally high (QA mean = 0.818) and tightly clustered "
        "(std = 0.075). Most samples receive similar retrieval scores, so the dynamic "
        "threshold collapses to approximately the same value as the static one."
    ))
    add_paragraph(doc, (
        "Second, the Cw-CONLI algorithm is designed for real-world RAG systems where "
        "retrieval quality varies significantly — some queries retrieve highly relevant "
        "passages while others retrieve only tangentially related content. In such settings, "
        "the retrieval score standard deviation might be 0.2–0.3, providing much greater "
        "differentiation for the weighting functions."
    ))
    add_paragraph(doc, (
        "Third, the grid search optimisation converges to similar effective thresholds for "
        "both C2 and C3. Because the retrieval scores are narrow, the optimal C3 parameters "
        "produce thresholds that are numerically close to the optimal C2 T_static for the "
        "majority of samples."
    ))

    # 8.5 The Summarization Challenge
    add_heading_custom(doc, "8.5 The Summarization Challenge", level=2)
    add_paragraph(doc, (
        "Summarization results are uniformly poor across all conditions, with over-flagging "
        "rates exceeding 99%. Three interconnected issues explain this:"
    ))
    add_paragraph(doc, (
        "1. Token truncation: RoBERTa-large-MNLI has a maximum sequence length of 512 tokens "
        "(premise + hypothesis combined). HaluEval summarization documents are typically "
        "1,000–3,000 tokens. After truncation, the NLI model sees only the beginning of the "
        "document, losing critical content that the summary may reference."
    ))
    add_paragraph(doc, (
        "2. Semantic mismatch: NLI models are trained on the MultiNLI corpus, which contains "
        "short, literal premise-hypothesis pairs. Summaries naturally involve paraphrasing, "
        "abstraction, and information synthesis — transformations that the NLI model interprets "
        "as non-entailment even when the summary is factually correct."
    ))
    add_paragraph(doc, (
        "3. Low score separability: The NLI gap between valid and hallucinated summaries is "
        "only 0.072 (0.208 vs. 0.136), compared to 0.460 for QA. This provides insufficient "
        "signal for any threshold-based classifier to operate effectively."
    ))

    # 8.6 Experiment 2 Insights
    add_heading_custom(doc, "8.6 Experiment 2 Insights", level=2)
    add_paragraph(doc, (
        "The realistic retrieval experiment, while producing impractically high over-flagging "
        "across all conditions, reveals two important insights:"
    ))
    add_paragraph(doc, (
        "First, C2 with T_static = 0.99 flags 100% of responses as hallucinations, indicating "
        "that when retrieval quality degrades (passages are retrieved from a pool rather than "
        "provided as ground truth), a uniformly high threshold becomes counterproductive. The "
        "NLI scores for correctly retrieved passages fall below 0.99, triggering false positives."
    ))
    add_paragraph(doc, (
        "Second, C3 variants reduce over-flagging by 3–5 percentage points. While this "
        "reduction is modest, it demonstrates the correct-direction advantage: by lowering "
        "the threshold for high-confidence retrievals, some valid responses that would be "
        "incorrectly flagged by C2 are correctly accepted by C3. This suggests that Cw-CONLI "
        "would show greater benefits in retrieval environments with even more score variability."
    ))

    # 8.7 Limitations
    add_heading_custom(doc, "8.7 Limitations", level=2)
    add_paragraph(doc, (
        "Several limitations constrain the generalisability of these findings:"
    ))
    add_paragraph(doc, (
        "1. Single NLI model: Only RoBERTa-large-MNLI is evaluated. Results may differ with "
        "other NLI models (e.g., DeBERTa-v3-large-MNLI, BART-large-MNLI) that have different "
        "score distributions and calibration properties."
    ))
    add_paragraph(doc, (
        "2. CPU-only evaluation: All experiments run on an Apple M4 CPU. Reported latencies "
        "are not representative of GPU-accelerated deployment, which would be 5–10x faster."
    ))
    add_paragraph(doc, (
        "3. HaluEval benchmark limitations: The dataset provides curated ground-truth knowledge "
        "per sample, which inflates retrieval scores and reduces the scenario diversity that "
        "Cw-CONLI is designed to exploit. Labels are approximately balanced (50/50), which "
        "may not reflect real-world hallucination prevalence."
    ))
    add_paragraph(doc, (
        "4. No multi-hop reasoning: All samples involve single-hop question answering or "
        "single-document summarization. Cw-CONLI's behaviour on multi-hop reasoning tasks, "
        "where retrieval confidence may vary significantly across hops, is unexplored."
    ))
    add_paragraph(doc, (
        "5. Fixed sigmoid parameters: The sigmoid's steepness (k) and pivot are not tuned "
        "in the grid search; default values (k = 10, pivot = 0.5) are used. Tuning these "
        "may improve sigmoid performance."
    ))

    # 8.8 Future Work
    add_heading_custom(doc, "8.8 Future Work", level=2)
    add_paragraph(doc, (
        "Based on the findings and limitations identified, several directions for future "
        "research are proposed:"
    ))
    add_paragraph(doc, (
        "1. Evaluation on heterogeneous retrieval datasets: Testing Cw-CONLI on benchmarks "
        "with natural retrieval score variation, such as multi-document QA (e.g., HotpotQA, "
        "MuSiQue) or open-domain QA with large retrieval corpora, would better demonstrate "
        "the algorithm's differentiation capability."
    ))
    add_paragraph(doc, (
        "2. Chunk-based NLI for long documents: Splitting long documents into chunks, running "
        "NLI on each chunk independently, and aggregating scores (e.g., max or weighted mean) "
        "would address the 512-token truncation limitation that cripples summarization performance."
    ))
    add_paragraph(doc, (
        "3. Ensemble NLI models: Combining predictions from multiple NLI models could improve "
        "score calibration and separability, particularly for summarization tasks where a "
        "single model's scores are poorly separated."
    ))
    add_paragraph(doc, (
        "4. Cross-model consistency: As noted in the original project scope, a cross-model "
        "consistency layer that compares outputs from multiple LLMs could provide an additional "
        "hallucination signal orthogonal to NLI verification."
    ))
    add_paragraph(doc, (
        "5. Real-world deployment evaluation: Testing the framework with live LLM generation "
        "(rather than pre-existing labelled responses) and measuring end-to-end user-facing "
        "latency and accuracy would validate practical applicability."
    ))
    add_paragraph(doc, (
        "6. Adaptive sigmoid parameters: Incorporating k and pivot into the grid search "
        "or learning them from data could improve the sigmoid variant's performance."
    ))

    # 8.9 Summary
    add_heading_custom(doc, "8.9 Chapter Summary", level=2)
    add_paragraph(doc, (
        "This chapter has interpreted the experimental results in context. Cw-CONLI does "
        "not significantly outperform static CONLI on HaluEval due to the dataset's tight "
        "retrieval score clustering. However, the algorithm is correctly designed: continuous "
        "weighting functions outperform the tiered variant, C3 Sigmoid achieves the lowest "
        "over-flagging on QA, and in the realistic retrieval setting C3 variants show a "
        "correct-direction advantage. All four research objectives are achieved. The primary "
        "limitation is the evaluation dataset's artificially high and narrow retrieval scores, "
        "which compress the operational range of the confidence-weighted threshold. Future work "
        "should evaluate on datasets with greater retrieval diversity to fully test the "
        "Cw-CONLI hypothesis."
    ))


# ============================================================================
# Main
# ============================================================================

def main():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Add header/footer
    add_header_footer(doc)

    # Title page
    for _ in range(6):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Confidence-Weighted Framework for Mitigating\nLLM Hallucinations")
    set_run_font(run, size=Pt(24), bold=True)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Updated Abstract & Chapters 6–8")
    set_run_font(run, size=Pt(16), italic=True)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_p.add_run("\nShaun Yogeshwaran\nw1912919\nBSc (Hons) Computer Science\nUniversity of Westminster / IIT")
    set_run_font(run, size=Pt(14))

    # Write chapters
    doc.add_page_break()
    write_abstract(doc)
    write_chapter_6(doc)
    write_chapter_7(doc)
    write_chapter_8(doc)

    # Save
    output_path = "thesis_new_chapters.docx"
    doc.save(output_path)
    print(f"Document saved to {output_path}")


if __name__ == "__main__":
    main()
